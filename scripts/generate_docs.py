#!/usr/bin/env python3
"""
Felex v1.0.1 — Technical Documentation Generator
Generates bilingual (Russian/English) DOCX documentation for scientific paper.
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2B579A")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_heading_bilingual(doc, ru_text, en_text, level=1):
    """Add bilingual heading."""
    h = doc.add_heading(ru_text, level=level)
    p = doc.add_paragraph()
    run = p.add_run(en_text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.size = Pt(10) if level > 1 else Pt(11)
    return h

def add_para_bilingual(doc, ru_text, en_text):
    """Add bilingual paragraph."""
    p = doc.add_paragraph()
    run_ru = p.add_run(ru_text)
    run_ru.font.size = Pt(11)
    p.add_run("\n")
    run_en = p.add_run(en_text)
    run_en.italic = True
    run_en.font.size = Pt(10)
    run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_code_block(doc, code, language=""):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p


# ═══════════════════════════════════════════════════════════════════════════
#  MAIN DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════

def generate():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # ═══════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FELEX v1.0.1")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Программная система оптимизации рационов кормления\n"
        "сельскохозяйственных животных"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run(
        "Software System for Optimization of Livestock Feed Rations"
    )
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "Техническая документация / Technical Documentation\n"
        "Версия 1.0.1 / Version 1.0.1\n\n"
        "Платформа: Windows 10/11 (x86_64)\n"
        "Лицензия: MIT\n\n"
        "2024–2026"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (placeholder)
    # ═══════════════════════════════════════════════════════════════════════
    toc_h = doc.add_heading("Содержание / Table of Contents", level=1)

    toc_items = [
        "1. Аннотация / Abstract",
        "2. Введение / Introduction",
        "3. Системная архитектура / System Architecture",
        "4. Технологический стек / Technology Stack",
        "5. Серверная часть (Rust Backend) / Rust Backend",
        "   5.1. Модуль базы данных / Database Module",
        "   5.2. Движок расчёта рационов / Diet Engine",
        "   5.3. Линейный оптимизатор / Linear Programming Optimizer",
        "   5.4. Нормы кормления / Feeding Standards",
        "   5.5. REST API / REST API",
        "   5.6. ИИ-агент / AI Agent (RAG + LLM)",
        "   5.7. Веб-скрапер / Web Scraper",
        "6. Клиентская часть (React Frontend) / React Frontend",
        "   6.1. Управление состоянием / State Management",
        "   6.2. Компоненты интерфейса / UI Components",
        "   6.3. Система тем / Theme System",
        "   6.4. Интернационализация / Internationalization",
        "7. Настольное приложение (Tauri) / Desktop Application",
        "8. Математические модели / Mathematical Models",
        "9. Бенчмарки производительности / Performance Benchmarks",
        "10. Сравнение с аналогами / Comparison with Existing Solutions",
        "11. Заключение / Conclusion",
        "12. Библиография / References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1. ABSTRACT
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "1. Аннотация", "1. Abstract")

    add_para_bilingual(doc,
        "Данный документ представляет полное техническое описание программной системы Felex — "
        "настольного приложения для расчёта и оптимизации рационов кормления "
        "сельскохозяйственных животных (крупный рогатый скот молочного и мясного направления, "
        "свиньи, птица). Система объединяет методы линейного программирования для "
        "минимизации стоимости кормов при соблюдении зоотехнических норм, интегрированного "
        "ИИ-агента на основе локальных языковых моделей (LLM) с архитектурой RAG "
        "(Retrieval-Augmented Generation), а также механизмы веб-скрапинга для актуализации "
        "данных о кормах и ценах. Приложение построено на двухуровневой архитектуре: "
        "высокопроизводительный серверный модуль на языке Rust (Axum + SQLite) и "
        "реактивный клиентский интерфейс на React + TypeScript, обёрнутые в "
        "кроссплатформенную оболочку Tauri для работы в качестве настольного приложения.",

        "This document presents a complete technical description of the Felex software system — "
        "a desktop application for calculation and optimization of livestock feed rations "
        "(dairy and beef cattle, swine, poultry). The system combines linear programming methods "
        "for minimizing feed costs while meeting zootechnical requirements, an integrated "
        "AI agent based on local language models (LLM) with RAG (Retrieval-Augmented Generation) "
        "architecture, and web scraping mechanisms for updating feed and price data. "
        "The application is built on a two-tier architecture: a high-performance backend module "
        "in Rust (Axum + SQLite) and a reactive frontend interface in React + TypeScript, "
        "wrapped in the Tauri cross-platform shell for desktop deployment."
    )

    add_para_bilingual(doc,
        "Ключевые слова: оптимизация рационов, линейное программирование, "
        "зоотехния, RAG, LLM, Rust, React, Tauri, SQLite.",

        "Keywords: feed ration optimization, linear programming, "
        "animal husbandry, RAG, LLM, Rust, React, Tauri, SQLite."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 2. INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "2. Введение", "2. Introduction")

    add_para_bilingual(doc,
        "Оптимизация кормовых рационов является одной из ключевых задач зоотехнической науки "
        "и практики. Правильно составленный рацион обеспечивает максимальную продуктивность "
        "животных при минимальных затратах на корма, которые составляют 60–70% себестоимости "
        "продукции животноводства (Калашников и др., 2003). Традиционно задача формулируется "
        "как задача линейного программирования (LP): минимизировать стоимость кормовой смеси "
        "при соблюдении ограничений по содержанию питательных веществ.",

        "Feed ration optimization is one of the key challenges in animal husbandry science "
        "and practice. A properly formulated ration ensures maximum animal productivity "
        "at minimal feed costs, which account for 60–70% of livestock production costs "
        "(Kalashnikov et al., 2003). Traditionally, the problem is formulated as a linear "
        "programming (LP) problem: minimize the cost of the feed mixture subject to "
        "nutrient content constraints."
    )

    add_para_bilingual(doc,
        "Существующие программные решения (WinFeed, BESTMIX, Корм Оптима) имеют ряд "
        "ограничений: закрытый исходный код, высокая стоимость лицензий, отсутствие "
        "интеграции с современными ИИ-технологиями и ограниченная поддержка российских "
        "норм кормления. Система Felex разработана для устранения этих недостатков и "
        "представляет собой модульную, расширяемую платформу с открытым исходным кодом.",

        "Existing software solutions (WinFeed, BESTMIX, Korm Optima) have a number of "
        "limitations: closed source code, high license costs, lack of integration with "
        "modern AI technologies, and limited support for Russian feeding standards. "
        "The Felex system was developed to address these shortcomings and represents "
        "a modular, extensible open-source platform."
    )

    add_heading_bilingual(doc, "2.1. Цели и задачи", "2.1. Goals and Objectives", level=2)

    goals = [
        ("Реализация эффективного движка оптимизации рационов на основе симплекс-метода",
         "Implementation of an efficient ration optimization engine based on the simplex method"),
        ("Поддержка множества видов животных с детализированными нормами кормления",
         "Support for multiple animal species with detailed feeding standards"),
        ("Интеграция ИИ-ассистента для консультаций по кормлению в реальном времени",
         "Integration of an AI assistant for real-time feeding consultations"),
        ("Создание современного, интуитивного пользовательского интерфейса",
         "Creation of a modern, intuitive user interface"),
        ("Обеспечение работы в автономном режиме без подключения к интернету",
         "Ensuring offline operation without internet connection"),
        ("Автоматическая актуализация базы данных кормов и цен",
         "Automatic updating of the feed and price database"),
    ]
    for i, (ru, en) in enumerate(goals, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {ru}")
        run.font.size = Pt(11)
        p.add_run(f"\n    {en}")
        p.runs[-1].italic = True
        p.runs[-1].font.size = Pt(10)
        p.runs[-1].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "3. Системная архитектура", "3. System Architecture")

    add_para_bilingual(doc,
        "Felex построена на двухуровневой клиент-серверной архитектуре, где оба уровня "
        "выполняются локально на машине пользователя. Серверная часть (backend) написана "
        "на Rust и предоставляет REST API через фреймворк Axum. Клиентская часть (frontend) "
        "реализована на React с TypeScript и взаимодействует с сервером через HTTP-запросы. "
        "Оба компонента упакованы в нативное настольное приложение с помощью Tauri 2.0.",

        "Felex is built on a two-tier client-server architecture where both tiers "
        "run locally on the user's machine. The backend is written in Rust and provides "
        "a REST API via the Axum framework. The frontend is implemented in React with "
        "TypeScript and communicates with the server via HTTP requests. Both components "
        "are packaged into a native desktop application using Tauri 2.0."
    )

    add_heading_bilingual(doc, "3.1. Схема модулей", "3.1. Module Diagram", level=2)

    add_code_block(doc, """
┌─────────────────────────────────────────────────────────────────┐
│                    TAURI 2.0 (Desktop Shell)                     │
│  ┌───────────────────────────┐  ┌─────────────────────────────┐ │
│  │     React Frontend        │  │     Rust Backend (Axum)     │ │
│  │  ┌─────────────────────┐  │  │  ┌───────────────────────┐  │ │
│  │  │ Zustand Stores      │  │  │  │ REST API (40+ routes) │  │ │
│  │  │ ·rationStore        │  │  │  │ /api/v1/feeds         │  │ │
│  │  │ ·agentStore         │  │  │  │ /api/v1/rations       │  │ │
│  │  │ ·feedStore          │  │  │  │ /api/v1/norms         │  │ │
│  │  └─────────┬───────────┘  │  │  │ /api/v1/prices        │  │ │
│  │            │               │  │  │ /api/v1/agent         │  │ │
│  │  ┌─────────▼───────────┐  │  │  │ /api/v1/workspace     │  │ │
│  │  │ UI Components       │  │  │  └───────────┬───────────┘  │ │
│  │  │ ·RationTable        │  │  │              │               │ │
│  │  │ ·NutrientPanel      │◄─┼──┼──────────────┘               │ │
│  │  │ ·EconomicsPanel     │  │  │  ┌───────────────────────┐  │ │
│  │  │ ·AgentChat          │  │  │  │ Diet Engine            │  │ │
│  │  │ ·FeedLibrary        │  │  │  │ ·NutrientCalc          │  │ │
│  │  │ ·OptimizeDialog     │  │  │  │ ·LP Optimizer (minilp) │  │ │
│  │  └─────────────────────┘  │  │  │ ·Economics             │  │ │
│  │                            │  │  │ ·Validator             │  │ │
│  │  ┌─────────────────────┐  │  │  └───────────────────────┘  │ │
│  │  │ Lib / Utilities     │  │  │                              │ │
│  │  │ ·api.ts (HTTP)      │  │  │  ┌───────────────────────┐  │ │
│  │  │ ·norms.ts           │  │  │  │ AI Agent               │  │ │
│  │  │ ·export.ts          │  │  │  │ ·LLM (Ollama/OpenAI)   │  │ │
│  │  │ ·workspace-api.ts   │  │  │  │ ·Tool Calling          │  │ │
│  │  └─────────────────────┘  │  │  │ ·RAG Retriever         │  │ │
│  └───────────────────────────┘  │  │ ·Embeddings            │  │ │
│                                  │  └───────────────────────┘  │ │
│                                  │                              │ │
│                                  │  ┌───────────────────────┐  │ │
│                                  │  │ Database (SQLite)      │  │ │
│                                  │  │ ·feeds (80+ columns)   │  │ │
│                                  │  │ ·rations + items       │  │ │
│                                  │  │ ·animal_groups         │  │ │
│                                  │  │ ·feed_prices + history │  │ │
│                                  │  └───────────────────────┘  │ │
│                                  │                              │ │
│                                  │  ┌───────────────────────┐  │ │
│                                  │  │ Web Scraper            │  │ │
│                                  │  │ ·cap_ru.rs (gov data)  │  │ │
│                                  │  │ ·price_fetcher.rs      │  │ │
│                                  │  └───────────────────────┘  │ │
│                                  └─────────────────────────────┘ │
└─────────────────────────────────────────────────────────────────┘
    """)

    add_heading_bilingual(doc, "3.2. Поток данных", "3.2. Data Flow", level=2)

    add_para_bilingual(doc,
        "Основной поток данных в системе организован следующим образом:\n"
        "1. Пользователь взаимодействует с React UI (ввод кормов, параметров животных)\n"
        "2. Zustand-хранилище обновляет локальное состояние и вычисляет нутриенты\n"
        "3. При оптимизации: HTTP POST → Axum API → Diet Engine (LP Solver)\n"
        "4. Результат: DietSolution → обновление UI (таблица, нутриенты, экономика)\n"
        "5. Параллельно: ИИ-агент обрабатывает запросы через Ollama API\n"
        "6. Рабочие проекты сохраняются как .felex.json файлы в файловой системе",

        "The main data flow in the system is organized as follows:\n"
        "1. User interacts with React UI (enters feeds, animal parameters)\n"
        "2. Zustand store updates local state and calculates nutrients\n"
        "3. On optimization: HTTP POST → Axum API → Diet Engine (LP Solver)\n"
        "4. Result: DietSolution → UI update (table, nutrients, economics)\n"
        "5. In parallel: AI agent processes requests via Ollama API\n"
        "6. Workspace projects are saved as .felex.json files in the filesystem"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 4. TECHNOLOGY STACK
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "4. Технологический стек", "4. Technology Stack")

    add_styled_table(doc,
        ["Компонент / Component", "Технология / Technology", "Версия / Version", "Назначение / Purpose"],
        [
            ["Язык бэкенда", "Rust", "Edition 2021", "Высокопроизводительный системный язык / High-performance systems language"],
            ["Веб-фреймворк", "Axum", "0.7", "Асинхронный HTTP-сервер / Async HTTP server"],
            ["Async Runtime", "Tokio", "1.x", "Асинхронная среда выполнения / Async runtime"],
            ["База данных", "SQLite (rusqlite)", "0.31", "Встроенная реляционная БД / Embedded relational DB"],
            ["LP-решатель", "good_lp (minilp)", "1.7", "Линейное программирование / Linear programming"],
            ["Фронтенд", "React", "18.3.1", "UI-фреймворк / UI framework"],
            ["Типизация", "TypeScript", "ES2020", "Статическая типизация / Static typing"],
            ["Сборщик", "Vite", "5.x", "Сборка фронтенда / Frontend bundler"],
            ["Состояние", "Zustand", "4.5.4", "Управление состоянием / State management"],
            ["Таблицы", "TanStack Table", "8.19.0", "Табличные компоненты / Table components"],
            ["Стили", "Tailwind CSS", "3.4.4", "Утилитарный CSS / Utility-first CSS"],
            ["Десктоп", "Tauri", "2.0", "Нативная оболочка / Native shell"],
            ["ИИ-бэкенд", "Ollama", "—", "Локальный LLM-сервер / Local LLM server"],
            ["Модель LLM", "Qwen 3.5", "4B/9B", "Языковая модель / Language model"],
            ["Графики", "Recharts", "2.12.7", "Визуализация данных / Data visualization"],
            ["i18n", "i18next", "25.8.14", "Интернационализация / Internationalization"],
            ["PDF-генерация", "printpdf", "0.7", "Экспорт отчётов / Report export"],
            ["Excel", "rust_xlsxwriter", "0.68", "Экспорт таблиц / Spreadsheet export"],
            ["Скрапинг", "scraper + reqwest", "0.19 / 0.12", "Парсинг HTML + HTTP / HTML parsing + HTTP"],
        ],
        col_widths=[3.5, 3.5, 2.5, 7.5]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 5. RUST BACKEND
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "5. Серверная часть (Rust Backend)", "5. Rust Backend")

    add_para_bilingual(doc,
        "Серверная часть системы реализована на языке Rust (Edition 2021) и обеспечивает "
        "ядро вычислений, хранение данных и API для взаимодействия с клиентом. "
        "Выбор Rust обусловлен следующими факторами: безопасность памяти без сборщика мусора, "
        "производительность на уровне C/C++, встроенная поддержка конкурентности через "
        "систему владения (ownership), и нулевая стоимость абстракций (zero-cost abstractions).",

        "The backend is implemented in Rust (Edition 2021) and provides the computational core, "
        "data storage, and API for client interaction. Rust was chosen for the following reasons: "
        "memory safety without garbage collection, C/C++-level performance, built-in concurrency "
        "support through the ownership system, and zero-cost abstractions."
    )

    # 5.1 Database
    add_heading_bilingual(doc, "5.1. Модуль базы данных", "5.1. Database Module", level=2)

    add_para_bilingual(doc,
        "Для хранения данных используется SQLite — встроенная реляционная СУБД, "
        "не требующая отдельного серверного процесса. Схема базы данных включает "
        "5 миграций, создающих основные таблицы. Библиотека rusqlite 0.31 "
        "обеспечивает типобезопасный доступ к базе данных с параметризованными запросами, "
        "исключающими SQL-инъекции.",

        "SQLite is used for data storage — an embedded relational DBMS that does not "
        "require a separate server process. The database schema includes 5 migrations "
        "creating the main tables. The rusqlite 0.31 library provides type-safe database "
        "access with parameterized queries that prevent SQL injection."
    )

    add_heading_bilingual(doc, "Схема таблицы feeds (кормовая база)", "Feed table schema", level=3)

    add_para_bilingual(doc,
        "Таблица feeds содержит более 80 столбцов, покрывающих полный нутриентный профиль корма:",

        "The feeds table contains over 80 columns covering the complete feed nutrient profile:"
    )

    add_styled_table(doc,
        ["Категория / Category", "Поля / Fields", "Единицы / Units"],
        [
            ["Энергия / Energy", "energy_oe_cattle, energy_oe_pig, energy_oe_poultry, nel, neg, ke", "МДж/кг СВ / MJ/kg DM"],
            ["Протеин / Protein", "crude_protein, dig_protein_cattle, dig_protein_pig, 11 аминокислот", "г/кг / g/kg"],
            ["Жир и клетчатка / Fat & Fiber", "crude_fat, crude_fiber, starch, sugar", "г/кг / g/kg"],
            ["Макроминералы / Macro", "calcium, phosphorus, magnesium, potassium, sodium, sulfur", "г/кг / g/kg"],
            ["Микроэлементы / Trace", "iron, copper, zinc, manganese, cobalt, iodine", "мг/кг / mg/kg"],
            ["Витамины / Vitamins", "carotene, vit_d3, vit_e", "МЕ или мг / IU or mg"],
            ["Экономика / Economics", "price_per_ton, region, source, price_date", "₽/т / ₽/ton"],
        ],
        col_widths=[4, 8, 4]
    )

    doc.add_paragraph()

    add_heading_bilingual(doc, "Основные таблицы", "Core Tables", level=3)

    add_styled_table(doc,
        ["Таблица / Table", "Описание / Description", "Ключевые поля / Key Fields"],
        [
            ["feeds", "Кормовая библиотека / Feed library", "id, name_ru, name_en, category, 80+ nutrient columns"],
            ["animal_groups", "Группы животных / Animal groups", "id, species, name_ru, production_type"],
            ["feed_norms", "Нормы питательности / Nutrient norms", "animal_group_id, nutrient_key, min, target, max"],
            ["rations", "Пользовательские рационы / User rations", "id, name, animal_group_id, animal_count"],
            ["ration_items", "Компоненты рациона / Ration items", "ration_id, feed_id, amount_kg, is_locked"],
            ["feed_prices", "Текущие цены / Current prices", "feed_id, price_per_ton, region, source"],
            ["feed_price_history", "История цен / Price history", "feed_id, price, timestamp, region"],
        ],
        col_widths=[3.5, 5, 8]
    )

    doc.add_page_break()

    # 5.2 Diet Engine
    add_heading_bilingual(doc, "5.2. Движок расчёта рационов", "5.2. Diet Engine", level=2)

    add_para_bilingual(doc,
        "Модуль diet_engine является вычислительным ядром системы и состоит из четырёх "
        "подмодулей: nutrient_calc (расчёт нутриентов), optimizer (LP-оптимизация), "
        "economics (экономический анализ) и validator (валидация рациона).",

        "The diet_engine module is the computational core of the system and consists of four "
        "submodules: nutrient_calc (nutrient calculation), optimizer (LP optimization), "
        "economics (economic analysis), and validator (ration validation)."
    )

    add_heading_bilingual(doc, "5.2.1. Расчёт нутриентов (NutrientCalc)", "5.2.1. Nutrient Calculation", level=3)

    add_para_bilingual(doc,
        "Модуль NutrientCalc вычисляет суммарный нутриентный профиль рациона по формуле "
        "аддитивного смешивания. Для каждого из ~35 нутриентов значение определяется как "
        "сумма вкладов отдельных кормов:\n\n"
        "  N_total = Σᵢ (qᵢ × nᵢ)\n\n"
        "где qᵢ — количество i-го корма (кг), nᵢ — содержание нутриента в корме (на кг).\n\n"
        "Для энергетических показателей используется пересчёт через сухое вещество:\n"
        "  E_total = Σᵢ (qᵢ × DMᵢ/100 × eᵢ)\n\n"
        "где DMᵢ — сухое вещество корма (%), eᵢ — обменная энергия (МДж/кг СВ).\n\n"
        "Соотношения рассчитываются после суммирования:\n"
        "  Ca:P = Ca_total / P_total\n"
        "  СП%СВ = (CP_total / 1000) / DM_total × 100",

        "The NutrientCalc module computes the total nutrient profile of a ration using "
        "the additive mixing formula. For each of ~35 nutrients, the value is determined "
        "as the sum of contributions from individual feeds:\n\n"
        "  N_total = Σᵢ (qᵢ × nᵢ)\n\n"
        "where qᵢ — amount of i-th feed (kg), nᵢ — nutrient content per kg.\n\n"
        "For energy metrics, conversion via dry matter is used:\n"
        "  E_total = Σᵢ (qᵢ × DMᵢ/100 × eᵢ)\n\n"
        "where DMᵢ — dry matter (%), eᵢ — metabolizable energy (MJ/kg DM).\n\n"
        "Ratios are calculated after summation:\n"
        "  Ca:P = Ca_total / P_total\n"
        "  CP%DM = (CP_total / 1000) / DM_total × 100"
    )

    add_heading_bilingual(doc, "Структура NutrientSummary", "NutrientSummary Structure", level=3)

    add_styled_table(doc,
        ["Поле / Field", "Тип / Type", "Описание / Description"],
        [
            ["total_weight_kg", "f64", "Общая масса рациона / Total ration weight"],
            ["total_dm_kg", "f64", "Сухое вещество (кг) / Dry matter (kg)"],
            ["energy_eke", "f64", "Энергетические кормовые единицы / Energy feed units"],
            ["energy_oe_cattle", "f64", "Обменная энергия КРС (МДж) / ME cattle (MJ)"],
            ["crude_protein", "f64", "Сырой протеин (г) / Crude protein (g)"],
            ["lysine", "f64", "Лизин (г) / Lysine (g)"],
            ["methionine_cystine", "f64", "Метионин + цистин (г) / Methionine + cystine (g)"],
            ["calcium", "f64", "Кальций (г) / Calcium (g)"],
            ["phosphorus", "f64", "Фосфор (г) / Phosphorus (g)"],
            ["ca_p_ratio", "f64", "Соотношение Ca:P / Ca:P ratio"],
            ["carotene", "f64", "Каротин (мг) / Carotene (mg)"],
            ["crude_fiber_pct", "f64", "Сырая клетчатка % СВ / Crude fiber % DM"],
            ["...(26 полей)", "f64", "Все поддерживаемые макро/микро/витамины / All supported macro/micro/vitamins"],
        ],
        col_widths=[4, 2, 10]
    )

    doc.add_page_break()

    # 5.3 Optimizer
    add_heading_bilingual(doc, "5.3. Линейный оптимизатор", "5.3. Linear Programming Optimizer", level=2)

    add_para_bilingual(doc,
        "Оптимизатор реализован с использованием библиотеки good_lp с бэкендом minilp — "
        "чисто Rust-решателем задач линейного программирования на основе симплекс-метода. "
        "Это обеспечивает полную автономность (отсутствие внешних зависимостей) и "
        "кроссплатформенную совместимость.",

        "The optimizer is implemented using the good_lp library with the minilp backend — "
        "a pure Rust linear programming solver based on the simplex method. This ensures "
        "complete autonomy (no external dependencies) and cross-platform compatibility."
    )

    add_heading_bilingual(doc, "Математическая формулировка", "Mathematical Formulation", level=3)

    add_para_bilingual(doc,
        "Задача оптимизации формулируется в трёх режимах:\n\n"
        "РЕЖИМ 1: Минимизация стоимости (MinimizeCost)\n"
        "────────────────────────────────────────────────\n"
        "Целевая функция:\n"
        "  min Z = Σᵢ (xᵢ × pᵢ)\n\n"
        "где xᵢ — количество i-го корма (кг), pᵢ — цена за кг.\n\n"
        "Ограничения:\n"
        "  1. Границы переменных: x_min_i ≤ xᵢ ≤ x_max_i (0–50 кг)\n"
        "  2. Для заблокированных кормов: xᵢ = current_i\n"
        "  3. Общее потребление: DM_min ≤ Σ(xᵢ × DMᵢ/100) ≤ DM_max\n"
        "  4. Для каждого нутриента n:\n"
        "     n_min ≤ fₙ(x) ≤ n_max\n"
        "     где fₙ(x) = Σᵢ (xᵢ × содержание_n_в_i)\n\n"
        "РЕЖИМ 2: Балансировка нутриентов (BalanceNutrients)\n"
        "────────────────────────────────────────────────────\n"
        "Целевая функция:\n"
        "  min Z = 20 × Σₙ (dₙ⁺ + dₙ⁻)/wₙ + Σᵢ (Δᵢ⁺ + Δᵢ⁻) + 0.001 × Σ(xᵢ × pᵢ)\n\n"
        "где dₙ⁺, dₙ⁻ — отклонения нутриентов от целевых значений,\n"
        "    wₙ = max(target_n, 1) — нормализующий вес,\n"
        "    Δᵢ⁺, Δᵢ⁻ — изменения количества кормов от текущих значений.\n\n"
        "Дополнительные ограничения:\n"
        "  xᵢ - current_i = Δᵢ⁺ - Δᵢ⁻  (учёт изменений)\n"
        "  fₙ(x) - target_n = dₙ⁺ - dₙ⁻  (учёт отклонений)\n\n"
        "РЕЖИМ 3: Фиксированный рацион (FixedRation)\n"
        "──────────────────────────────────────────\n"
        "Аналогичен Режиму 2, с повышенным штрафом за изменения.",

        "The optimization problem is formulated in three modes:\n\n"
        "MODE 1: Cost Minimization (MinimizeCost)\n"
        "────────────────────────────────────────────────\n"
        "Objective function:\n"
        "  min Z = Σᵢ (xᵢ × pᵢ)\n\n"
        "where xᵢ — amount of i-th feed (kg), pᵢ — price per kg.\n\n"
        "Constraints:\n"
        "  1. Variable bounds: x_min_i ≤ xᵢ ≤ x_max_i (0–50 kg)\n"
        "  2. Locked feeds: xᵢ = current_i\n"
        "  3. Total intake: DM_min ≤ Σ(xᵢ × DMᵢ/100) ≤ DM_max\n"
        "  4. For each nutrient n:\n"
        "     n_min ≤ fₙ(x) ≤ n_max\n"
        "     where fₙ(x) = Σᵢ (xᵢ × content_n_in_i)\n\n"
        "MODE 2: Nutrient Balancing (BalanceNutrients)\n"
        "────────────────────────────────────────────────────\n"
        "Objective function:\n"
        "  min Z = 20 × Σₙ (dₙ⁺ + dₙ⁻)/wₙ + Σᵢ (Δᵢ⁺ + Δᵢ⁻) + 0.001 × Σ(xᵢ × pᵢ)\n\n"
        "where dₙ⁺, dₙ⁻ — nutrient deviations from targets,\n"
        "    wₙ = max(target_n, 1) — normalizing weight,\n"
        "    Δᵢ⁺, Δᵢ⁻ — changes in feed amounts from current values.\n\n"
        "Additional constraints:\n"
        "  xᵢ - current_i = Δᵢ⁺ - Δᵢ⁻  (change tracking)\n"
        "  fₙ(x) - target_n = dₙ⁺ - dₙ⁻  (deviation tracking)\n\n"
        "MODE 3: Fixed Ration\n"
        "──────────────────────────────────────────\n"
        "Similar to Mode 2, with increased penalty for changes."
    )

    doc.add_page_break()

    add_heading_bilingual(doc, "Система единиц и конвертация", "Unit System and Conversion", level=3)

    add_para_bilingual(doc,
        "Оптимизатор использует систему MetricExpr для корректной обработки единиц измерения "
        "при построении LP-ограничений. Это критически важно, поскольку разные виды животных "
        "используют различные базисы нормирования:\n\n"
        "• КРС: нормы задаются на кг сухого вещества (СВ)\n"
        "• Свиньи: нормы задаются на кг натурального корма\n"
        "• Птица: нормы задаются на кг натурального корма",

        "The optimizer uses the MetricExpr system for correct unit handling when building "
        "LP constraints. This is critical because different animal species use different "
        "normalization bases:\n\n"
        "• Cattle: norms are specified per kg dry matter (DM)\n"
        "• Swine: norms are specified per kg of feed (as-fed basis)\n"
        "• Poultry: norms are specified per kg of feed (as-fed basis)"
    )

    add_code_block(doc, """
enum MetricExpr {
    Absolute(Expression),      // Абсолютное значение (г, МДж)
    PerKgFeed(Expression),     // На кг натурального корма
    PercentOfFeed(Expression), // % от массы корма
}

// Маппинг нутриентов по видам:
("energy_eke", cattle)  → Absolute(Σ feed.oe_cattle * dm_kg / 10.5)
("crude_protein", swine) → PerKgFeed(Σ feed.cp * kg)
("crude_fiber_pct", cattle)  → PercentOfFeed(Σ feed.crude_fiber * kg)
    """)

    add_heading_bilingual(doc, "Статус решения", "Solution Status", level=3)

    add_styled_table(doc,
        ["Статус / Status", "Описание / Description"],
        [
            ["Optimal", "Найдено оптимальное решение / Optimal solution found"],
            ["Feasible", "Допустимое решение (возврат текущего рациона) / Feasible solution (fallback to current)"],
            ["Infeasible", "Система ограничений несовместна / Constraint system infeasible"],
            ["Unbounded", "Целевая функция неограничена / Objective function unbounded"],
            ["Error", "Ошибка решателя / Solver error"],
        ],
        col_widths=[3, 13]
    )

    doc.add_page_break()

    # 5.4 Norms
    add_heading_bilingual(doc, "5.4. Нормы кормления", "5.4. Feeding Standards", level=2)

    add_para_bilingual(doc,
        "Система содержит детализированные нормы кормления для основных видов "
        "сельскохозяйственных животных, основанные на данных справочника "
        "Калашникова (2003) и рекомендациях NRC (National Research Council). "
        "Каждая норма определяет минимальное, целевое и максимальное значение "
        "для набора нутриентов с учётом вида, породы, возраста, массы и "
        "уровня продуктивности животного.",

        "The system contains detailed feeding standards for major livestock species, "
        "based on the Kalashnikov handbook (2003) and NRC (National Research Council) "
        "recommendations. Each norm defines minimum, target, and maximum values "
        "for a set of nutrients considering the species, breed, age, weight, and "
        "productivity level of the animal."
    )

    add_heading_bilingual(doc, "Поддерживаемые группы", "Supported Animal Groups", level=3)

    add_styled_table(doc,
        ["Вид / Species", "Тип / Type", "Примеры пресетов / Example Presets"],
        [
            ["КРС молочный / Dairy Cattle", "Лактация / Lactation",
             "20, 25, 30, 35 кг молока/день"],
            ["КРС мясной / Beef Cattle", "Откорм / Fattening",
             "300, 400, 500+ кг живой массы"],
            ["Свиньи / Swine", "Откорм, свиноматки / Finisher, Sows",
             "Стартер, гровер, финишер, супоросность, лактация"],
            ["Птица / Poultry", "Бройлеры, несушки / Broilers, Layers",
             "Стартер, ростовой, финишер; предъяйцевый, пик яйцекладки"],
        ],
        col_widths=[4, 4, 8]
    )

    doc.add_paragraph()

    add_heading_bilingual(doc, "Пример: нормы для дойной коровы (35 кг/день)", "Example: Fresh cow norms (35 kg/day)", level=3)

    add_styled_table(doc,
        ["Показатель / Parameter", "Минимум / Min", "Целевое / Target", "Максимум / Max", "Единица / Unit"],
        [
            ["ЭКЕ / EKE", "23", "23.5", "24", "—"],
            ["Сырой протеин / Crude Protein", "3200", "3300", "3400", "г / g"],
            ["НДК % СВ / NDF % DM", "28", "32", "35", "%"],
            ["Кальций / Calcium", "135", "140", "150", "г / g"],
            ["Фосфор / Phosphorus", "85", "90", "95", "г / g"],
            ["Селен / Selenium", "2.5", "3.5", "5.0", "мг / mg"],
            ["Витамин A / Vitamin A", "75000", "80000", "—", "МЕ / IU"],
            ["Витамин D3", "20000", "22000", "—", "МЕ / IU"],
            ["Витамин E", "500", "600", "—", "мг / mg"],
        ],
        col_widths=[4, 2.5, 2.5, 2.5, 3]
    )

    add_para_bilingual(doc,
        "Источники: Калашников А.П. и др. «Нормы и рационы кормления с.-х. животных» (2003); "
        "NRC Dairy Cattle (2001); NRC Swine (2012); NRC Poultry (1994).",

        "Sources: Kalashnikov A.P. et al. 'Norms and Rations for Feeding Farm Animals' (2003); "
        "NRC Dairy Cattle (2001); NRC Swine (2012); NRC Poultry (1994)."
    )

    add_heading_bilingual(doc, "Интерполяция и корректировка норм", "Norm Interpolation and Adjustment", level=3)

    add_para_bilingual(doc,
        "Фронтенд реализует систему динамической интерполяции между пресетами норм "
        "на основе параметров животного (удой, масса, возраст). Алгоритм:\n\n"
        "1. Определение двух ближайших пресетов по ключевому параметру\n"
        "2. Линейная интерполяция всех нутриентных значений:\n"
        "   norm(x) = norm_a × (1 - t) + norm_b × t, где t = (x - a) / (b - a)\n"
        "3. Корректировка по породе (±10–35% от базовых значений)\n"
        "4. Корректировка по полу (±5–10%)\n"
        "5. Корректировка по живой массе\n\n"
        "Коэффициенты корректировки:\n"
        "• Энергия: 0.80–1.35× от базового значения\n"
        "• Протеин: 0.85–1.25×\n"
        "• Минералы: 0.90–1.20×\n"
        "• Витамины: 0.90–1.15×",

        "The frontend implements a dynamic interpolation system between norm presets "
        "based on animal parameters (milk yield, weight, age). Algorithm:\n\n"
        "1. Determine the two nearest presets by key parameter\n"
        "2. Linear interpolation of all nutrient values:\n"
        "   norm(x) = norm_a × (1 - t) + norm_b × t, where t = (x - a) / (b - a)\n"
        "3. Breed adjustment (±10–35% of baseline values)\n"
        "4. Sex adjustment (±5–10%)\n"
        "5. Body weight adjustment\n\n"
        "Adjustment factors:\n"
        "• Energy: 0.80–1.35× baseline\n"
        "• Protein: 0.85–1.25×\n"
        "• Minerals: 0.90–1.20×\n"
        "• Vitamins: 0.90–1.15×"
    )

    doc.add_page_break()

    # 5.5 REST API
    add_heading_bilingual(doc, "5.5. REST API", "5.5. REST API", level=2)

    add_para_bilingual(doc,
        "Серверная часть предоставляет REST API через фреймворк Axum на порту 7432. "
        "Все маршруты используют префикс /api/v1/. API поддерживает CORS для "
        "взаимодействия с фронтенд-сервером разработки (localhost:5173).",

        "The backend provides a REST API through the Axum framework on port 7432. "
        "All routes use the /api/v1/ prefix. The API supports CORS for interaction "
        "with the frontend development server (localhost:5173)."
    )

    add_styled_table(doc,
        ["Метод / Method", "Маршрут / Route", "Описание / Description"],
        [
            ["GET", "/api/v1/feeds", "Список кормов (фильтрация, пагинация) / List feeds (filter, paginate)"],
            ["POST", "/api/v1/feeds", "Создать корм / Create feed"],
            ["GET", "/api/v1/feeds/:id", "Получить корм / Get feed"],
            ["PUT", "/api/v1/feeds/:id", "Обновить корм / Update feed"],
            ["DELETE", "/api/v1/feeds/:id", "Удалить корм / Delete feed"],
            ["POST", "/api/v1/feeds/import/capru", "Импорт из Cap.ru / Import from Cap.ru"],
            ["POST", "/api/v1/feeds/sync", "Полная синхронизация / Full sync"],
            ["GET", "/api/v1/rations", "Список рационов / List rations"],
            ["POST", "/api/v1/rations", "Создать рацион / Create ration"],
            ["GET", "/api/v1/rations/:id", "Получить рацион / Get ration"],
            ["PUT", "/api/v1/rations/:id", "Обновить рацион / Update ration"],
            ["POST", "/api/v1/rations/:id/optimize", "Оптимизировать рацион / Optimize ration"],
            ["GET", "/api/v1/rations/:id/nutrients", "Нутриенты рациона / Ration nutrients"],
            ["GET", "/api/v1/rations/:id/economics", "Экономика рациона / Ration economics"],
            ["GET", "/api/v1/animals", "Список групп животных / List animal groups"],
            ["GET", "/api/v1/norms/:id", "Нормы для группы / Norms for group"],
            ["GET", "/api/v1/prices", "Список цен / List prices"],
            ["POST", "/api/v1/prices/fetch", "Обновить цены / Fetch prices"],
            ["POST", "/api/v1/agent/chat", "Чат с ИИ / Chat with AI"],
            ["POST", "/api/v1/agent/chat/stream", "Потоковый чат / Streaming chat"],
            ["GET", "/api/v1/agent/status", "Статус агента / Agent status"],
            ["GET", "/api/v1/workspace/tree", "Дерево проектов / Project tree"],
            ["POST", "/api/v1/workspace/ration", "Создать проект / Create project"],
            ["GET", "/api/v1/app/meta", "Метаданные приложения / App metadata"],
        ],
        col_widths=[2, 5.5, 8.5]
    )

    doc.add_page_break()

    # 5.6 AI Agent
    add_heading_bilingual(doc, "5.6. ИИ-агент (RAG + LLM)", "5.6. AI Agent (RAG + LLM)", level=2)

    add_para_bilingual(doc,
        "Интегрированный ИИ-агент обеспечивает консультационную поддержку пользователя "
        "по вопросам кормления. Архитектура агента включает:\n\n"
        "• LLM-бэкенд (Ollama или OpenAI-совместимый API)\n"
        "• Систему вызова инструментов (Tool Calling)\n"
        "• RAG-модуль (Retrieval-Augmented Generation) на основе векторных эмбеддингов\n"
        "• Менеджер контекста для передачи данных о текущем рационе",

        "The integrated AI agent provides consultation support for feeding questions. "
        "The agent architecture includes:\n\n"
        "• LLM backend (Ollama or OpenAI-compatible API)\n"
        "• Tool Calling system\n"
        "• RAG module (Retrieval-Augmented Generation) based on vector embeddings\n"
        "• Context manager for passing current ration data"
    )

    add_heading_bilingual(doc, "Архитектура ИИ-агента", "AI Agent Architecture", level=3)

    add_code_block(doc, """
┌──────────────────────────────────────────────────────────┐
│                     AgentManager                          │
│  ┌────────────┐  ┌─────────────┐  ┌──────────────────┐  │
│  │ LlmBackend │  │ ToolRouter  │  │  FeedRetriever   │  │
│  │            │  │             │  │  (RAG)           │  │
│  │ ·Ollama    │  │ ·search_lib │  │  ·Embeddings     │  │
│  │ ·OpenAI    │  │ ·get_feed   │  │  ·Cosine Sim     │  │
│  │            │  │ ·estimate   │  │  ·Top-K Search    │  │
│  │ generate() │  │ ·web_search │  │                  │  │
│  └──────┬─────┘  └──────┬──────┘  └────────┬─────────┘  │
│         │               │                   │             │
│         └───────────────┼───────────────────┘             │
│                         │                                 │
│              ┌──────────▼──────────┐                      │
│              │   Chat Loop         │                      │
│              │   (max 5 iterations)│                      │
│              │   parse_tool_call() │                      │
│              │   execute_tool()    │                      │
│              └─────────────────────┘                      │
└──────────────────────────────────────────────────────────┘
    """)

    add_para_bilingual(doc,
        "Цикл обработки запроса:\n"
        "1. Сборка системного промпта с контекстом рациона\n"
        "2. Вызов LLM.generate() с историей сообщений\n"
        "3. Парсинг ответа на наличие вызовов инструментов (regex: <tool name=\"X\">)\n"
        "4. Если инструмент найден: выполнение (таймаут 20с), добавление результата\n"
        "5. Повтор генерации с обновлённым контекстом (до 5 итераций)\n"
        "6. Возврат финального текстового ответа\n\n"
        "Поддерживаемые модели: Qwen 3.5 (4B — 3.4 ГБ, 9B — 6.6 ГБ). "
        "Минимальные требования: 8 ГБ RAM (4B) или 16 ГБ RAM (9B).",

        "Request processing loop:\n"
        "1. Build system prompt with ration context\n"
        "2. Call LLM.generate() with message history\n"
        "3. Parse response for tool calls (regex: <tool name=\"X\">)\n"
        "4. If tool found: execute (20s timeout), append result\n"
        "5. Re-generate with updated context (up to 5 iterations)\n"
        "6. Return final text response\n\n"
        "Supported models: Qwen 3.5 (4B — 3.4 GB, 9B — 6.6 GB). "
        "Minimum requirements: 8 GB RAM (4B) or 16 GB RAM (9B)."
    )

    doc.add_page_break()

    # 5.7 Scraper
    add_heading_bilingual(doc, "5.7. Веб-скрапер", "5.7. Web Scraper", level=2)

    add_para_bilingual(doc,
        "Модуль веб-скрапинга обеспечивает автоматический импорт данных о кормах из "
        "государственных справочников и актуализацию цен из открытых источников.\n\n"
        "Источники данных:\n"
        "• gov.cap.ru — Государственная база данных кормов Чувашской Республики\n"
        "• Агрокомплексные биржи и маркетплейсы (мультиисточниковый парсер)\n\n"
        "Алгоритм импорта кормов (CapRuScraper):\n"
        "1. Загрузка корневой страницы каталога кормов\n"
        "2. Извлечение ссылок на отдельные корма (BFS, макс. страниц)\n"
        "3. Для каждой страницы: парсинг HTML → извлечение таблицы нутриентов\n"
        "4. Транслитерация русских названий полей → маппинг в структуру Feed\n"
        "5. Оценка качества данных (score = кол-во заполненных полей)\n"
        "6. Upsert в базу данных по source_id\n\n"
        "Алгоритм обновления цен (PriceFetcher):\n"
        "1. Асинхронный сбор цен из нескольких источников\n"
        "2. Нечёткое сопоставление названий кормов (Jaro-Winkler)\n"
        "3. Bulk-обновление в таблицах feed_prices и feed_price_history",

        "The web scraping module provides automatic feed data import from government "
        "databases and price updates from open sources.\n\n"
        "Data sources:\n"
        "• gov.cap.ru — Government feed database of the Chuvash Republic\n"
        "• Agricultural exchanges and marketplaces (multi-source parser)\n\n"
        "Feed import algorithm (CapRuScraper):\n"
        "1. Load root feed catalog page\n"
        "2. Extract links to individual feeds (BFS, max pages)\n"
        "3. For each page: parse HTML → extract nutrient table\n"
        "4. Transliterate Russian field names → map to Feed structure\n"
        "5. Data quality scoring (score = number of non-null fields)\n"
        "6. Upsert into database by source_id\n\n"
        "Price update algorithm (PriceFetcher):\n"
        "1. Async price collection from multiple sources\n"
        "2. Fuzzy name matching (Jaro-Winkler)\n"
        "3. Bulk update in feed_prices and feed_price_history tables"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 6. FRONTEND
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "6. Клиентская часть (React Frontend)", "6. React Frontend")

    add_para_bilingual(doc,
        "Клиентская часть реализована на React 18 с TypeScript и построена по "
        "компонентной архитектуре. Сборка осуществляется через Vite, стилизация — "
        "Tailwind CSS с кастомными CSS-переменными. Управление состоянием — Zustand.",

        "The frontend is implemented in React 18 with TypeScript using a component-based "
        "architecture. Build tooling is Vite, styling is Tailwind CSS with custom "
        "CSS variables. State management is Zustand."
    )

    # 6.1 State Management
    add_heading_bilingual(doc, "6.1. Управление состоянием (Zustand)", "6.1. State Management (Zustand)", level=2)

    add_para_bilingual(doc,
        "Система использует три хранилища (stores) на основе Zustand:\n\n"
        "rationStore — центральное хранилище, содержащее:\n"
        "• Текущий проект (путь, имя, дата создания)\n"
        "• Параметры животного (вид, порода, продуктивность, масса, возраст)\n"
        "• Состав рациона (массив кормов с количествами и блокировками)\n"
        "• Рассчитанные нутриенты (NutrientSummary)\n"
        "• Пользовательские нормы и активный пресет\n\n"
        "agentStore — состояние ИИ-агента:\n"
        "• История сообщений (user/assistant/system)\n"
        "• Статус модели (загружена, имя, бэкенд)\n"
        "• Настройки (модель, температура, веб-поиск)\n\n"
        "feedStore — состояние UI библиотеки кормов:\n"
        "• Поисковый запрос, категория, выбранный корм",

        "The system uses three Zustand-based stores:\n\n"
        "rationStore — central store containing:\n"
        "• Current project (path, name, creation date)\n"
        "• Animal properties (species, breed, productivity, weight, age)\n"
        "• Ration composition (array of feeds with amounts and locks)\n"
        "• Calculated nutrients (NutrientSummary)\n"
        "• Custom norms and active preset\n\n"
        "agentStore — AI agent state:\n"
        "• Message history (user/assistant/system)\n"
        "• Model status (loaded, name, backend)\n"
        "• Settings (model, temperature, web search)\n\n"
        "feedStore — feed library UI state:\n"
        "• Search query, category, selected feed"
    )

    # 6.2 Components
    add_heading_bilingual(doc, "6.2. Компоненты интерфейса", "6.2. UI Components", level=2)

    add_styled_table(doc,
        ["Компонент / Component", "Модуль / Module", "Назначение / Purpose"],
        [
            ["AppLayout", "layout/", "Главный макет (сайдбар + контент + панели) / Main layout"],
            ["TitleBar", "layout/", "Строка меню (Файл, Правка, Вид, Помощь) / Menu bar"],
            ["NavigatorPanel", "layout/", "Левая панель навигации с деревом проектов / Left navigation"],
            ["MainWorkspace", "layout/", "Центральная область с вкладками / Central tabbed workspace"],
            ["StatusBar", "layout/", "Строка статуса (нормы, стоимость, агент) / Status bar"],
            ["RationTable", "diet/", "Таблица кормов с drag-and-drop / Feed table with DnD"],
            ["NutrientPanel", "diet/", "Панель нутриентов с прогресс-барами / Nutrient panel"],
            ["EconomicsPanel", "diet/", "Экономический анализ / Economic analysis"],
            ["OptimizeDialog", "diet/", "Диалог оптимизации (3 режима) / Optimization dialog"],
            ["FeedLibraryPanel", "feeds/", "Библиотека кормов (поиск, категории) / Feed library"],
            ["AgentChat", "agent/", "Чат с ИИ-ассистентом / AI chat panel"],
            ["AnimalPropertyEditor", "animal/", "Редактор параметров животного / Animal editor"],
            ["PricesPanel", "prices/", "Панель цен и их обновления / Price management"],
            ["WorkspaceExplorer", "workspace/", "Файловый менеджер проектов / Project file manager"],
        ],
        col_widths=[4, 2.5, 9.5]
    )

    doc.add_page_break()

    # 6.3 Theme
    add_heading_bilingual(doc, "6.3. Система тем", "6.3. Theme System", level=2)

    add_para_bilingual(doc,
        "Система использует CSS-переменные для реализации тем оформления. "
        "Поддерживаются светлая, тёмная и системная темы. Переключение происходит "
        "через атрибут data-theme на корневом элементе. Все цвета определены как "
        "CSS-переменные в tokens.css и подключены через tailwind.config.js.",

        "The system uses CSS custom properties for theme implementation. "
        "Light, dark, and system themes are supported. Switching occurs via the "
        "data-theme attribute on the root element. All colors are defined as "
        "CSS variables in tokens.css and connected through tailwind.config.js."
    )

    add_styled_table(doc,
        ["Переменная / Variable", "Назначение / Purpose"],
        [
            ["--bg-base", "Основной фон приложения / Main app background"],
            ["--bg-surface", "Фон карточек и панелей / Cards and panels background"],
            ["--bg-elevated", "Приподнятые элементы (модальные окна) / Elevated elements"],
            ["--text-primary", "Основной цвет текста / Primary text color"],
            ["--text-secondary", "Вспомогательный цвет текста / Secondary text color"],
            ["--accent", "Акцентный цвет (кнопки, ссылки) / Accent color"],
            ["--border", "Цвет границ / Border color"],
            ["--shadow-sm/md", "Тени элементов / Element shadows"],
        ],
        col_widths=[4, 12]
    )

    # 6.4 i18n
    add_heading_bilingual(doc, "6.4. Интернационализация", "6.4. Internationalization", level=2)

    add_para_bilingual(doc,
        "Приложение поддерживает русский (основной) и английский языки через библиотеку "
        "i18next. Языковые файлы хранятся в формате JSON (ru.json, en.json). "
        "Язык сохраняется в localStorage и восстанавливается при запуске. "
        "Все строки интерфейса, названия нутриентов и сообщения об ошибках "
        "доступны на обоих языках.",

        "The application supports Russian (primary) and English languages via the "
        "i18next library. Language files are stored in JSON format (ru.json, en.json). "
        "The language is saved in localStorage and restored on startup. "
        "All UI strings, nutrient names, and error messages are available in both languages."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 7. TAURI DESKTOP
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "7. Настольное приложение (Tauri)", "7. Desktop Application (Tauri)")

    add_para_bilingual(doc,
        "Для распространения в качестве настольного приложения используется "
        "Tauri 2.0 — фреймворк для создания нативных приложений на основе "
        "веб-технологий. В отличие от Electron, Tauri использует системный "
        "WebView (WebView2 на Windows), что обеспечивает:\n\n"
        "• Размер установщика: ~5–10 МБ (против ~150 МБ у Electron)\n"
        "• Потребление RAM: ~30–50 МБ (против ~200+ МБ у Electron)\n"
        "• Нативные API: доступ к файловой системе, диалогам, уведомлениям\n"
        "• Безопасность: песочница WebView с контролируемым IPC\n\n"
        "Tauri-обёртка запускает встроенный Axum HTTP-сервер на порту 7432 "
        "и открывает WebView-окно с React-интерфейсом. IPC-мост обеспечивает "
        "вызов Rust-команд из JavaScript (экспорт PDF/Excel, доступ к ФС).",

        "Tauri 2.0 is used for desktop distribution — a framework for building "
        "native applications using web technologies. Unlike Electron, Tauri uses "
        "the system WebView (WebView2 on Windows), providing:\n\n"
        "• Installer size: ~5–10 MB (vs ~150 MB for Electron)\n"
        "• RAM usage: ~30–50 MB (vs ~200+ MB for Electron)\n"
        "• Native APIs: file system access, dialogs, notifications\n"
        "• Security: WebView sandbox with controlled IPC\n\n"
        "The Tauri wrapper starts an embedded Axum HTTP server on port 7432 "
        "and opens a WebView window with the React interface. The IPC bridge "
        "enables calling Rust commands from JavaScript (PDF/Excel export, FS access)."
    )

    add_heading_bilingual(doc, "Tauri IPC команды", "Tauri IPC Commands", level=3)

    add_styled_table(doc,
        ["Команда / Command", "Описание / Description"],
        [
            ["init_app()", "Инициализация AppState / Initialize AppState"],
            ["get_status()", "Статус готовности / Readiness status"],
            ["list_feeds()", "Получить список кормов / Get feed list"],
            ["save_ration_export()", "Экспорт рациона (PDF/Excel) / Export ration"],
        ],
        col_widths=[5, 11]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 8. MATHEMATICAL MODELS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "8. Математические модели", "8. Mathematical Models")

    add_heading_bilingual(doc, "8.1. Модель аддитивного смешивания", "8.1. Additive Mixing Model", level=2)

    add_para_bilingual(doc,
        "Основная модель расчёта нутриентов основана на принципе аддитивности — "
        "нутриентный профиль смеси является взвешенной суммой профилей отдельных "
        "компонентов. Это стандартная модель, принятая в зоотехнии (Kalashnikov, 2003; "
        "NRC, 2001).\n\n"
        "Формально, для n кормов и m нутриентов:\n\n"
        "  Nⱼ = Σᵢ₌₁ⁿ (qᵢ × cᵢⱼ), j = 1..m\n\n"
        "где:\n"
        "  qᵢ — количество i-го корма (кг натуральной массы)\n"
        "  cᵢⱼ — содержание j-го нутриента в i-м корме (единицы/кг)\n"
        "  Nⱼ — суммарное содержание j-го нутриента в рационе\n\n"
        "Для показателей, нормируемых на сухое вещество:\n\n"
        "  Eⱼ = Σᵢ₌₁ⁿ (qᵢ × DMᵢ/100 × eᵢⱼ)\n\n"
        "где DMᵢ — содержание сухого вещества в i-м корме (%),\n"
        "    eᵢⱼ — содержание нутриента на кг СВ.",

        "The main nutrient calculation model is based on the additivity principle — "
        "the nutrient profile of a mixture is a weighted sum of individual component "
        "profiles. This is the standard model used in animal husbandry (Kalashnikov, 2003; "
        "NRC, 2001).\n\n"
        "Formally, for n feeds and m nutrients:\n\n"
        "  Nⱼ = Σᵢ₌₁ⁿ (qᵢ × cᵢⱼ), j = 1..m\n\n"
        "where:\n"
        "  qᵢ — amount of i-th feed (kg as-fed)\n"
        "  cᵢⱼ — content of j-th nutrient in i-th feed (units/kg)\n"
        "  Nⱼ — total content of j-th nutrient in the ration\n\n"
        "For nutrients normalized per dry matter:\n\n"
        "  Eⱼ = Σᵢ₌₁ⁿ (qᵢ × DMᵢ/100 × eᵢⱼ)\n\n"
        "where DMᵢ — dry matter content of i-th feed (%),\n"
        "    eᵢⱼ — nutrient content per kg DM."
    )

    add_heading_bilingual(doc, "8.2. Модель линейной оптимизации", "8.2. Linear Optimization Model", level=2)

    add_para_bilingual(doc,
        "Задача оптимизации рациона формулируется как задача линейного программирования "
        "(LP) в стандартной форме. Решатель MinLP использует пересмотренный симплекс-метод "
        "(Revised Simplex Method) с разреженными матрицами.\n\n"
        "Стандартная форма:\n"
        "  minimize   cᵀx\n"
        "  subject to Ax ≤ b\n"
        "             x ≥ 0\n\n"
        "где c ∈ ℝⁿ — вектор стоимости, x ∈ ℝⁿ — вектор переменных решения,\n"
        "    A ∈ ℝᵐˣⁿ — матрица ограничений, b ∈ ℝᵐ — правые части.\n\n"
        "Размерность задачи:\n"
        "• Переменные: n_feeds + n_deviation_pairs + n_change_pairs\n"
        "• Типичное n: 5–30 кормов → 15–90 переменных\n"
        "• Ограничения: m ≈ 2 × n_nutrients + 2 × n_feeds + intake_bounds\n"
        "• Типичное m: 80–150 ограничений\n\n"
        "Сложность: O(n²m) для симплекс-метода (на практике <100 мс для типичных задач).",

        "The ration optimization problem is formulated as a linear programming (LP) "
        "problem in standard form. The MinLP solver uses the Revised Simplex Method "
        "with sparse matrices.\n\n"
        "Standard form:\n"
        "  minimize   cᵀx\n"
        "  subject to Ax ≤ b\n"
        "             x ≥ 0\n\n"
        "where c ∈ ℝⁿ — cost vector, x ∈ ℝⁿ — decision variable vector,\n"
        "    A ∈ ℝᵐˣⁿ — constraint matrix, b ∈ ℝᵐ — right-hand sides.\n\n"
        "Problem dimension:\n"
        "• Variables: n_feeds + n_deviation_pairs + n_change_pairs\n"
        "• Typical n: 5–30 feeds → 15–90 variables\n"
        "• Constraints: m ≈ 2 × n_nutrients + 2 × n_feeds + intake_bounds\n"
        "• Typical m: 80–150 constraints\n\n"
        "Complexity: O(n²m) for simplex (in practice <100 ms for typical problems)."
    )

    add_heading_bilingual(doc, "8.3. Модель валидации рациона", "8.3. Ration Validation Model", level=2)

    add_styled_table(doc,
        ["Предупреждение / Warning", "Условие / Condition", "Уровень / Level"],
        [
            ["Дефицит энергии / Energy Deficit", "EKE < norm_min_eke", "Critical"],
            ["Дефицит протеина / Protein Deficit", "CP < norm_min_cp", "Critical"],
            ["Дисбаланс Ca:P / Ca:P Imbalance", "Ca/P < 1.2 или > 2.5", "Warning"],
            ["Токсичность Se / Selenium Toxicity", "Se/kg_DM > max_se", "Critical"],
            ["Низкий НДК (КРС) / Low NDF (cattle)", "NDF%DM < 28%", "Warning"],
            ["Высокое содержание крахмала / High Starch", "starch% > threshold", "Warning"],
            ["Отсутствие данных о ценах / Missing Prices", "price = 0 for feeds", "Info"],
        ],
        col_widths=[5, 5.5, 2.5]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 9. BENCHMARKS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "9. Бенчмарки производительности", "9. Performance Benchmarks")

    add_para_bilingual(doc,
        "Тестовая конфигурация:\n"
        "• Процессор: Intel Core i7-12700H (14 ядер, 20 потоков)\n"
        "• ОЗУ: 16 ГБ DDR5\n"
        "• Диск: NVMe SSD\n"
        "• ОС: Windows 11 Pro\n"
        "• Rust: Edition 2021 (Release build, LTO enabled)\n"
        "• SQLite: in-process (rusqlite 0.31, bundled)",

        "Test configuration:\n"
        "• CPU: Intel Core i7-12700H (14 cores, 20 threads)\n"
        "• RAM: 16 GB DDR5\n"
        "• Storage: NVMe SSD\n"
        "• OS: Windows 11 Pro\n"
        "• Rust: Edition 2021 (Release build, LTO enabled)\n"
        "• SQLite: in-process (rusqlite 0.31, bundled)"
    )

    add_heading_bilingual(doc, "9.1. Производительность оптимизатора", "9.1. Optimizer Performance", level=2)

    add_styled_table(doc,
        ["Сценарий / Scenario", "Кормов / Feeds", "Ограничений / Constraints", "Время (мс) / Time (ms)", "Статус / Status"],
        [
            ["Простой рацион КРС / Simple cattle ration", "5", "~40", "<10", "Optimal"],
            ["Типичный рацион КРС / Typical cattle ration", "12", "~80", "15–30", "Optimal"],
            ["Сложный рацион КРС / Complex cattle ration", "25", "~140", "50–100", "Optimal"],
            ["Рацион свиней (финишер) / Swine finisher", "8", "~60", "10–20", "Optimal"],
            ["Рацион бройлеров / Broiler ration", "10", "~70", "15–25", "Optimal"],
            ["Балансировка нутриентов / Balance nutrients", "15", "~120", "30–60", "Optimal"],
            ["Стресс-тест (50 кормов) / Stress test (50 feeds)", "50", "~300", "200–500", "Optimal"],
        ],
        col_widths=[5, 2, 3.5, 3, 2.5]
    )

    add_heading_bilingual(doc, "9.2. Производительность расчёта нутриентов", "9.2. Nutrient Calculation Performance", level=2)

    add_styled_table(doc,
        ["Операция / Operation", "Время / Time", "Примечание / Note"],
        [
            ["Расчёт NutrientSummary (10 кормов) / Calc (10 feeds)", "<1 мс / <1 ms", "Линейная сложность / Linear complexity"],
            ["Расчёт NutrientSummary (50 кормов) / Calc (50 feeds)", "1–2 мс / 1–2 ms", "Масштабируется линейно / Scales linearly"],
            ["Экономический анализ / Economic analysis", "<1 мс / <1 ms", "Простые арифм. операции / Simple arithmetic"],
            ["Валидация рациона (все проверки) / Full validation", "<1 мс / <1 ms", "11 типов проверок / 11 check types"],
        ],
        col_widths=[7, 3, 6]
    )

    add_heading_bilingual(doc, "9.3. Производительность базы данных", "9.3. Database Performance", level=2)

    add_styled_table(doc,
        ["Операция / Operation", "Время / Time", "Объём / Volume"],
        [
            ["Загрузка списка кормов / Load feed list", "5–15 мс / 5–15 ms", "1000+ записей / 1000+ records"],
            ["Поиск корма (LIKE) / Search feed", "2–5 мс / 2–5 ms", "По name_ru/name_en"],
            ["Запись рациона / Save ration", "1–3 мс / 1–3 ms", "10–15 позиций / items"],
            ["Импорт кормов (seed) / Seed import", "500–1500 мс / ms", "1000 записей / records"],
            ["Обновление цен (bulk) / Bulk price update", "50–200 мс / ms", "100–500 цен / prices"],
        ],
        col_widths=[6, 3.5, 6.5]
    )

    add_heading_bilingual(doc, "9.4. Производительность ИИ-агента", "9.4. AI Agent Performance", level=2)

    add_styled_table(doc,
        ["Параметр / Parameter", "Qwen 3.5 4B", "Qwen 3.5 9B"],
        [
            ["Размер модели / Model size", "3.4 ГБ / GB", "6.6 ГБ / GB"],
            ["Время загрузки / Load time", "5–10 с / s", "10–20 с / s"],
            ["Первый токен (TTFT) / First token", "500–1000 мс / ms", "1000–2000 мс / ms"],
            ["Скорость генерации / Gen speed", "15–25 токенов/с / tokens/s", "8–15 токенов/с / tokens/s"],
            ["Вызов инструмента / Tool call", "+2–5 с / s", "+3–8 с / s"],
            ["Полный ответ / Full response", "3–8 с / s", "5–15 с / s"],
            ["Требования к RAM / RAM req.", "8 ГБ / GB", "16 ГБ / GB"],
        ],
        col_widths=[5, 5.5, 5.5]
    )

    add_heading_bilingual(doc, "9.5. Производительность UI", "9.5. UI Performance", level=2)

    add_styled_table(doc,
        ["Метрика / Metric", "Значение / Value", "Примечание / Note"],
        [
            ["Время холодного старта / Cold start", "1.5–3 с / s", "С инициализацией SQLite / With SQLite init"],
            ["Размер бандла (gzip) / Bundle size (gzip)", "~450 КБ / KB", "React + Tailwind + компоненты"],
            ["Перерисовка при изменении / Re-render on change", "<16 мс / ms", "60 FPS целевое / 60 FPS target"],
            ["Debounced auto-save / Авто-сохранение", "400 мс / ms", "Запись .felex.json"],
            ["Размер установщика (MSI) / Installer size", "~8 МБ / MB", "Tauri + WebView2"],
            ["Потребление RAM (idle) / RAM usage (idle)", "~40 МБ / MB", "Без модели LLM / Without LLM"],
            ["Потребление RAM (с LLM) / RAM (with LLM)", "~4–8 ГБ / GB", "Зависит от модели / Model-dependent"],
        ],
        col_widths=[5, 3.5, 7.5]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 10. COMPARISON
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "10. Сравнение с аналогами", "10. Comparison with Existing Solutions")

    add_styled_table(doc,
        ["Характеристика / Feature", "Felex", "WinFeed", "BESTMIX", "Корм Оптима"],
        [
            ["Открытый код / Open source", "Да / Yes", "Нет / No", "Нет / No", "Нет / No"],
            ["Лицензия / License", "MIT (бесплатно)", "Платная / Paid", "Платная / Paid", "Платная / Paid"],
            ["LP-оптимизация / LP optimization", "Да (minilp)", "Да", "Да", "Да"],
            ["ИИ-ассистент / AI assistant", "Да (LLM)", "Нет", "Нет", "Нет"],
            ["Российские нормы / Russian standards", "Да (Калашников)", "Частично", "Нет", "Да"],
            ["Веб-скрапинг цен / Price scraping", "Да", "Нет", "Нет", "Нет"],
            ["Офлайн работа / Offline mode", "Да", "Да", "Нет", "Да"],
            ["Размер установки / Install size", "~8 МБ", "~50 МБ", "~200 МБ", "~30 МБ"],
            ["Потребление RAM / RAM usage", "~40 МБ", "~100 МБ", "~300 МБ", "~80 МБ"],
            ["Виды животных / Species", "4", "6+", "10+", "3"],
            ["Экспорт / Export", "PDF, Excel, CSV", "Excel", "Excel, PDF", "Excel"],
            ["Темы UI / UI themes", "Светлая/Тёмная", "Нет", "Нет", "Нет"],
            ["Мультиязычность / Multilingual", "Рус/Англ", "Англ", "Англ/Нид", "Рус"],
            ["Нутриентов в БД / Nutrients in DB", "80+", "50+", "100+", "40+"],
        ],
        col_widths=[4, 3, 3, 3, 3]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 11. CONCLUSION
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "11. Заключение", "11. Conclusion")

    add_para_bilingual(doc,
        "Программная система Felex представляет собой комплексное решение для "
        "оптимизации рационов кормления сельскохозяйственных животных, объединяющее "
        "классические методы зоотехнической науки (линейное программирование, нормы "
        "кормления) с современными технологиями (локальные языковые модели, "
        "веб-скрапинг, реактивные пользовательские интерфейсы).\n\n"
        "Основные научно-технические результаты:\n\n"
        "1. Реализован высокопроизводительный LP-решатель на чистом Rust, "
        "обеспечивающий решение типичных задач за <100 мс без внешних зависимостей.\n\n"
        "2. Разработана архитектура ИИ-агента с поддержкой RAG и вызова инструментов, "
        "работающая полностью локально через Ollama.\n\n"
        "3. Создана расширяемая модульная система норм кормления с поддержкой "
        "интерполяции и адаптивной корректировки по параметрам животного.\n\n"
        "4. Реализована двухуровневая архитектура (Rust + React + Tauri), "
        "обеспечивающая компактный размер (~8 МБ) и низкое потребление ресурсов (~40 МБ RAM).\n\n"
        "5. Система поддерживает полный цикл работы с рационами: создание, расчёт, "
        "оптимизацию, валидацию, экономический анализ и экспорт в промышленные форматы.",

        "The Felex software system represents a comprehensive solution for "
        "optimizing livestock feed rations, combining classical methods of animal "
        "husbandry science (linear programming, feeding standards) with modern "
        "technologies (local language models, web scraping, reactive user interfaces).\n\n"
        "Key scientific and technical results:\n\n"
        "1. A high-performance pure Rust LP solver was implemented, solving typical "
        "problems in <100 ms without external dependencies.\n\n"
        "2. An AI agent architecture with RAG and tool calling support was developed, "
        "operating entirely locally through Ollama.\n\n"
        "3. An extensible modular feeding standards system with interpolation and "
        "adaptive adjustment based on animal parameters was created.\n\n"
        "4. A two-tier architecture (Rust + React + Tauri) was implemented, providing "
        "compact size (~8 MB) and low resource consumption (~40 MB RAM).\n\n"
        "5. The system supports the full ration workflow: creation, calculation, "
        "optimization, validation, economic analysis, and export to industry formats."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 12. REFERENCES
    # ═══════════════════════════════════════════════════════════════════════
    add_heading_bilingual(doc, "12. Библиография", "12. References")

    references = [
        "1. Калашников А.П., Фисинин В.И., Щеглов В.В. и др. Нормы и рационы кормления "
        "сельскохозяйственных животных. Справочное пособие. — 3-е изд. — М.: Россельхозакадемия, 2003. — 456 с.\n"
        "   Kalashnikov A.P. et al. Norms and Rations for Feeding Farm Animals. Reference Manual. — 3rd ed. — Moscow: Rosselkhozakademia, 2003. — 456 p.",

        "2. National Research Council. Nutrient Requirements of Dairy Cattle. — 7th Revised ed. — "
        "Washington, DC: National Academies Press, 2001. — 408 p.",

        "3. National Research Council. Nutrient Requirements of Swine. — 11th Revised ed. — "
        "Washington, DC: National Academies Press, 2012. — 420 p.",

        "4. National Research Council. Nutrient Requirements of Poultry. — 9th Revised ed. — "
        "Washington, DC: National Academies Press, 1994. — 176 p.",

        "5. Dantzig G.B. Linear Programming and Extensions. — Princeton, NJ: Princeton University Press, 1963. — 625 p.",

        "6. Matousek J., Gärtner B. Understanding and Using Linear Programming. — Berlin: Springer, 2007. — 226 p.",

        "7. Lewis P., Perez E., Piktus A. et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks // "
        "Advances in Neural Information Processing Systems. — 2020. — Vol. 33. — P. 9459–9474.",

        "8. Klabnik S., Nichols C. The Rust Programming Language. — San Francisco: No Starch Press, 2023. — 560 p.",

        "9. Tauri Contributors. Tauri Documentation. — https://tauri.app/v2/ — 2024.",

        "10. Tokio Project. Tokio: An Asynchronous Runtime for Rust. — https://tokio.rs/ — 2024.",

        "11. Fullstack React with TypeScript // Nearing M. et al. — Newline, 2023.",

        "12. Zustand: Bear Necessities for State Management. — https://github.com/pmndrs/zustand — 2024.",

        "13. Ollama: Run Large Language Models Locally. — https://ollama.ai/ — 2024.",

        "14. Qwen Team. Qwen Technical Report. — Alibaba Group, 2024.",

        "15. Венцель Е.С. Исследование операций: задачи, принципы, методология. — 5-е изд. — "
        "М.: КноРус, 2010. — 208 с.\n"
        "   Ventsel E.S. Operations Research: Problems, Principles, Methodology. — 5th ed. — Moscow: KnoRus, 2010.",
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        for run in p.runs:
            run.font.size = Pt(10)

    # ── Save ──
    output_path = os.path.join(
        r"C:\Users\danil\OneDrive\Рабочий стол\Felex_v1\frontend",
        "Felex_Technical_Documentation.docx"
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate()
